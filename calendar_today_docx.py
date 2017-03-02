from __future__ import print_function
import httplib2
import os

from apiclient import discovery
from oauth2client import client
from oauth2client import tools
from oauth2client.file import Storage

from docx import Document

import datetime

try:
    import argparse
    flags = argparse.ArgumentParser(parents=[tools.argparser]).parse_args()
except ImportError:
    flags = None

# If modifying these scopes, delete your previously saved credentials
# at ~/.credentials/calendar-python-quickstart.json
SCOPES = 'https://www.googleapis.com/auth/calendar.readonly'
CLIENT_SECRET_FILE = 'C:\\Users\\externo101\\Documents\\Python\\calendar\\client_secret.json'
APPLICATION_NAME = 'Google Calendar API Python Quickstart'


def get_credentials():
    """Gets valid user credentials from storage.

    If nothing has been stored, or if the stored credentials are invalid,
    the OAuth2 flow is completed to obtain the new credentials.

    Returns:
        Credentials, the obtained credential.
    """
    home_dir = os.path.expanduser('~')
    credential_dir = os.path.join(home_dir, '.credentials')
    if not os.path.exists(credential_dir):
        os.makedirs(credential_dir)
    credential_path = os.path.join(credential_dir,
                                   'calendar-python-quickstart.json')

    store = Storage(credential_path)
    credentials = store.get()
    if not credentials or credentials.invalid:
        flow = client.flow_from_clientsecrets(CLIENT_SECRET_FILE, SCOPES)
        flow.user_agent = APPLICATION_NAME
        if flags:
            credentials = tools.run_flow(flow, store, flags)
        else: # Needed only for compatibility with Python 2.6
            credentials = tools.run(flow, store)
        print('Storing credentials to ' + credential_path)
    return credentials

def main():
    """Shows basic usage of the Google Calendar API.

    Creates a Google Calendar API service object and outputs a list of the next
    10 events on the user's calendar.
    """
    credentials = get_credentials()
    http = credentials.authorize(httplib2.Http())
    service = discovery.build('calendar', 'v3', http=http)

    ti = datetime.datetime.now()
    AVUI = ti.strftime('%Y-%m-%dT23:59:59Z')	#to End of today
    now = datetime.datetime.utcnow().isoformat() + 'Z' # 'Z' indicates UTC time
    print('Getting the upcoming 10 events')
    eventsResult = service.events().list(
        calendarId='primary', timeMin=now, timeMax=AVUI, maxResults=10, singleEvents=True,
        orderBy='startTime').execute()
    events = eventsResult.get('items', [])

    document = Document('demo2.docx')
    document.add_page_break()   #separació de la plantilla
    document.add_heading('Agenda', 0)

    if not events:
        print('No upcoming events found.')
    for event in events:
        start = event['start'].get('dateTime', event['start'].get('date'))
        print(start, event['summary'])
        document.add_paragraph(start[11:16] + " : " +event['summary'], style=None)

    document.save('demo3.docx')

if __name__ == '__main__':
    main()

#https://gist.github.com/cwurld/9b4e10dbeecab28345a3
